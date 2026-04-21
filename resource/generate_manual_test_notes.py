from __future__ import annotations

import argparse
from io import BytesIO
from pathlib import Path

from docx import Document


NOTES = {
    "logic.docx": [
        (
            "1. 命题逻辑与真值表",
            [
                "命题逻辑研究命题、逻辑联结词与推理规则。",
                "真值表可以用来判断命题公式在不同赋值下的真假。",
                "常见联结词包括非、且、或、蕴含和等价。",
            ],
        ),
        (
            "2. 等值演算与范式",
            [
                "利用等值演算可以把公式转化为主析取范式或主合取范式。",
                "命题公式是否为永真式或矛盾式，也可以通过真值表判断。",
            ],
        ),
    ],
    "relation.docx": [
        (
            "1. 二元关系",
            [
                "二元关系的基本性质包括自反性、对称性、反对称性和传递性。",
                "等价关系会把集合划分成若干等价类。",
            ],
        ),
        (
            "2. 偏序与哈斯图",
            [
                "偏序关系满足自反、反对称和传递。",
                "哈斯图适合表示集合包含关系和整除关系。",
            ],
        ),
    ],
    "graph.docx": [
        (
            "1. 图与连通性",
            [
                "图论研究顶点、边、路径、回路和连通分量。",
                "无向图中，顶点的度数反映与其相连的边数。",
            ],
        ),
        (
            "2. 欧拉回路与生成树",
            [
                "欧拉回路要求图连通且所有顶点度数均为偶数。",
                "生成树保持图连通，同时不含回路。",
            ],
        ),
    ],
}


def build_docx(path: Path, sections: list[tuple[str, list[str]]]) -> None:
    document = Document()
    for heading, paragraphs in sections:
        document.add_paragraph(heading)
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)

    buffer = BytesIO()
    document.save(buffer)
    path.write_bytes(buffer.getvalue())


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate mock DOCX notes for manual browser E2E testing.")
    parser.add_argument(
        "--output",
        default=".tmp/manual_browser_e2e_notes",
        help="Output directory for generated DOCX files.",
    )
    args = parser.parse_args()

    output_dir = Path(args.output)
    output_dir.mkdir(parents=True, exist_ok=True)

    for filename, sections in NOTES.items():
        build_docx(output_dir / filename, sections)

    print("Generated mock notes:")
    for filename in NOTES:
        print(" -", (output_dir / filename).as_posix())


if __name__ == "__main__":
    main()
