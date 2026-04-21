# -*- coding: utf-8 -*-
from __future__ import annotations

import shutil
import socket
import threading
import time
from io import BytesIO
from pathlib import Path

import requests
import uvicorn
from docx import Document
from playwright.sync_api import Page, expect, sync_playwright

import app.main as main_module
import app.routers.auth as auth_router
import app.routers.knowledge as knowledge_router
import app.routers.note as note_router
import app.routers.query as query_router
import app.routers.schedule as schedule_router
import app.services as services_module
import app.services.note_service as note_service_module
import app.storage.schedule_store as schedule_store_module
import app.storage.user_store as user_store_module
from app.config import HF_CACHE_DIR
from app.services.auth_service import AuthService
from app.services.knowledge_service import KnowledgeService
from app.services.note_service import NoteService
from app.services.schedule_service import ScheduleService
from app.services.topic_vector_store import TopicVectorStore
from app.storage.knowledge_tree_store import KnowledgeTreeStore
from app.storage.note_store import NoteStore
from app.storage.schedule_store import ScheduleStore
from app.storage.user_store import UserStore

REGISTER_NAME = "离散数学测试"
COURSE_NAME = "离散数学"
TEACHER_NAME = "周老师"
LOCATION_NAME = "理工楼B201"

TOPIC_LOGIC = "命题逻辑"
TOPIC_RELATION = "关系与偏序"
TOPIC_GRAPH = "图论"

QUERY_TRUTH_TABLE = "真值表"
QUERY_EULER = "欧拉回路"
GRAPH_ROUTED_TAG = "已主题路由"

BTN_SHOW_ALL = "查看全部"
BTN_ASSIGN = "关联到当前主题"
BTN_SEARCH = "检索切片"
BTN_REFRESH_GRAPH = "刷新图谱"


class IntegrationScraperStub:
    def fetch_schedule(self, account, password, semester_id, *, prefer_playwright=False):
        return []

    def parse_pdf_schedule(self, content: bytes):
        return []


def make_docx_file(path: Path, sections: list[tuple[str, list[str]]]) -> Path:
    document = Document()
    for heading, paragraphs in sections:
        document.add_paragraph(heading)
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)

    buffer = BytesIO()
    document.save(buffer)
    path.write_bytes(buffer.getvalue())
    return path


def reserve_port() -> int:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
        sock.bind(("127.0.0.1", 0))
        return int(sock.getsockname()[1])


def wait_for_server_ready(base_url: str, timeout_seconds: float = 20.0) -> None:
    deadline = time.time() + timeout_seconds
    last_error: Exception | None = None
    while time.time() < deadline:
        try:
            response = requests.get(base_url + "/login", timeout=2)
            if response.status_code == 200:
                return
        except Exception as exc:
            last_error = exc
        time.sleep(0.2)
    if last_error is not None:
        raise RuntimeError(f"Server did not start in time: {last_error}") from last_error
    raise RuntimeError("Server did not start in time")


def patch_note_file_dir(note_files_dir: Path) -> None:
    note_service_module.NOTE_FILES_DIR = note_files_dir
    NoteService.__init__.__globals__["NOTE_FILES_DIR"] = note_files_dir
    NoteService.upload.__globals__["NOTE_FILES_DIR"] = note_files_dir
    NoteService.delete.__globals__["NOTE_FILES_DIR"] = note_files_dir
    NoteService.get_file_path.__globals__["NOTE_FILES_DIR"] = note_files_dir


def configure_isolated_app(runtime_root: Path) -> Path:
    data_dir = runtime_root / "data"
    users_file = data_dir / "users.json"
    schedules_dir = data_dir / "schedules"
    logs_dir = data_dir / "logs"
    note_files_dir = data_dir / "note_files"
    notes_db_path = data_dir / "notes.db"
    knowledge_tree_dir = data_dir / "knowledge_trees"
    topic_store_dir = data_dir / "qdrant_db" / "topic_vectors"

    user_store_module.USERS_FILE = users_file
    schedule_store_module.SCHEDULES_DIR = schedules_dir
    patch_note_file_dir(note_files_dir)

    main_module.DATA_DIR = data_dir
    main_module.USERS_FILE = users_file
    main_module.SCHEDULES_DIR = schedules_dir
    main_module.LOGS_DIR = logs_dir
    main_module.APP_LOG_FILE = logs_dir / "app.log"
    main_module.NOTE_FILES_DIR = note_files_dir

    user_store = UserStore()
    schedule_store = ScheduleStore()
    note_store = NoteStore(db_path=notes_db_path)
    scraper = IntegrationScraperStub()
    auth_service = AuthService(user_store)
    schedule_service = ScheduleService(schedule_store, scraper)
    note_service = NoteService(note_store)
    knowledge_service = KnowledgeService(
        note_store,
        tree_store=KnowledgeTreeStore(root_dir=knowledge_tree_dir),
    )

    knowledge_service._memory_failed = True
    knowledge_service._llm_failed = True
    knowledge_service._topic_store_failed = False
    knowledge_service._topic_store = TopicVectorStore(
        root_dir=topic_store_dir,
        cache_dir=HF_CACHE_DIR,
    )
    knowledge_service._topic_store._ensure_client()

    services_module.user_store = user_store
    services_module.schedule_store = schedule_store
    services_module.note_store = note_store
    services_module.scnu_scraper = scraper
    services_module.auth_service = auth_service
    services_module.schedule_service = schedule_service
    services_module.note_service = note_service
    services_module.knowledge_service = knowledge_service
    main_module.auth_service = auth_service

    auth_router.auth_service = auth_service
    schedule_router.auth_service = auth_service
    schedule_router.schedule_service = schedule_service
    query_router.auth_service = auth_service
    query_router.schedule_service = schedule_service
    note_router.auth_service = auth_service
    note_router.note_service = note_service
    note_router.knowledge_service = knowledge_service
    knowledge_router.auth_service = auth_service
    knowledge_router.knowledge_service = knowledge_service

    return note_files_dir


def create_docx_fixtures(runtime_root: Path) -> dict[str, Path]:
    notes_dir = runtime_root / "fixtures"
    notes_dir.mkdir(parents=True, exist_ok=True)
    return {
        "logic": make_docx_file(
            notes_dir / "logic.docx",
            [
                (
                    "1. 命题逻辑与等值演算",
                    [
                        "命题逻辑研究命题、联结词、真值表与等值演算。真值表可以用来判断命题公式的真假。",
                    ],
                ),
                (
                    "2. 范式与推理",
                    [
                        "主析取范式和主合取范式用于化简命题公式。利用等值演算可以判断推理是否有效。",
                    ],
                ),
            ],
        ),
        "relation": make_docx_file(
            notes_dir / "relation.docx",
            [
                (
                    "1. 二元关系",
                    [
                        "关系的基本性质包括自反性、对称性、反对称性和传递性。等价关系会把集合划分为等价类。",
                    ],
                ),
                (
                    "2. 偏序与哈斯图",
                    [
                        "偏序关系满足自反、反对称和传递。哈斯图可以展示整除关系与集合包含关系。",
                    ],
                ),
            ],
        ),
        "graph": make_docx_file(
            notes_dir / "graph.docx",
            [
                (
                    "1. 图与连通性",
                    [
                        "图论研究顶点、边、路径与连通分量。无向图的度数和连通性用于分析结构。",
                    ],
                ),
                (
                    "2. 欧拉回路与生成树",
                    [
                        "欧拉回路要求图连通且所有顶点度数为偶数。生成树保持连通而不产生回路。",
                    ],
                ),
            ],
        ),
    }


def register_and_login(page: Page, base_url: str, student_id: str, password: str) -> None:
    page.goto(base_url + "/login", wait_until="networkidle")
    expect(page.locator("#auth-app")).to_be_visible()

    page.locator("#auth-app .auth-actions .button.secondary").click()

    text_inputs = page.locator("#auth-app input[type='text']")
    text_inputs.nth(0).fill(REGISTER_NAME)
    text_inputs.nth(1).fill(student_id)
    text_inputs.nth(2).fill(student_id)
    page.locator("#auth-app input[type='password']").fill(password)
    page.locator("#auth-app .auth-actions .button").nth(0).click()

    page.wait_for_url(base_url + "/dashboard", timeout=30_000)
    expect(page.locator("#dashboard-app")).to_be_visible()


def create_schedule_and_course(page: Page) -> None:
    side_panels = page.locator(".side-column section.panel")
    semester_panel = side_panels.nth(0)
    course_form_panel = side_panels.nth(2)

    semester_panel.locator("input[type='text']").nth(0).fill("2025-2026-2")
    semester_panel.locator("input[type='date']").fill("2026-03-02")
    semester_panel.locator("button").nth(0).click()

    course_form_panel.locator("input[type='text']").nth(0).fill(COURSE_NAME)
    course_form_panel.locator("input[type='text']").nth(1).fill(TEACHER_NAME)
    course_form_panel.locator("input[type='text']").nth(2).fill(LOCATION_NAME)
    course_form_panel.locator("select").nth(0).select_option("2")
    course_form_panel.locator("select").nth(1).select_option("all")
    course_form_panel.locator("input[type='number']").nth(0).fill("3")
    course_form_panel.locator("input[type='number']").nth(1).fill("4")
    course_form_panel.locator("input[type='text']").nth(3).fill("1-20")
    course_form_panel.locator("button[type='submit']").click()

    expect(page.locator(".course-list .course-item")).to_have_count(1)
    expect(page.locator(".timetable .course-card")).to_have_count(1)

    knowledge_links = page.locator("a[href*='/knowledge-workspace?course_id=']")
    assert knowledge_links.count() >= 2, "Expected at least two knowledge workspace entry links on dashboard"


def open_knowledge_workspace(page: Page) -> None:
    page.locator(".course-list .course-item a[href*='/knowledge-workspace?course_id=']").first.click()
    page.wait_for_url("**/knowledge-workspace?course_id=*", timeout=30_000)
    expect(page.locator("#knowledge-app")).to_be_visible()


def create_topic(page: Page, name: str, summary: str, keywords: str) -> None:
    form = page.locator(".topic-create-form")
    form.locator("input[type='text']").nth(0).fill(name)
    form.locator("textarea").fill(summary)
    form.locator("input[type='text']").nth(1).fill(keywords)
    form.locator("button[type='submit']").click()
    expect(page.locator(".tree-list")).to_contain_text(name)


def show_all_notes(page: Page) -> None:
    page.get_by_role("button", name=BTN_SHOW_ALL).click()


def upload_note(page: Page, file_path: Path, expected_filename: str) -> None:
    show_all_notes(page)
    before_count = page.locator(".note-card").count()
    page.locator(".upload-box input[type='file']").set_input_files(str(file_path))
    page.locator(".upload-bar button").click()
    expect(page.locator(".note-card")).to_have_count(before_count + 1, timeout=30_000)
    expect(page.locator(".note-list")).to_contain_text(expected_filename)


def assign_latest_note_to_topic(page: Page, topic_name: str, expected_filename: str) -> None:
    show_all_notes(page)
    page.locator(".note-card", has_text=expected_filename).click()
    page.locator(".tree-row", has_text=topic_name).click()
    assign_button = page.get_by_role("button", name=BTN_ASSIGN)
    if assign_button.count() > 0 and assign_button.is_visible():
        assign_button.click()
    expect(page.locator(".detail-tags")).to_contain_text(topic_name)
    expect(page.locator(".note-card.active .tag.accent")).to_contain_text(topic_name)


def verify_search_and_graph(page: Page) -> None:
    right_panel = page.locator(".knowledge-column-right section.panel").nth(1)
    text_inputs = right_panel.locator("input[type='text']")

    text_inputs.nth(0).fill(QUERY_TRUTH_TABLE)
    right_panel.get_by_role("button", name=BTN_SEARCH).click()
    first_result = right_panel.locator(".result-list .result-card").nth(0)
    expect(first_result).to_be_visible(timeout=30_000)
    expect(right_panel.locator(".result-list")).to_contain_text(QUERY_TRUTH_TABLE)

    text_inputs.nth(1).fill(QUERY_EULER)
    right_panel.get_by_role("button", name=BTN_REFRESH_GRAPH).click()
    expect(page.locator(".graph-tags")).to_contain_text(GRAPH_ROUTED_TAG, timeout=30_000)
    expect(page.locator(".graph-tags")).to_contain_text(TOPIC_GRAPH, timeout=30_000)


def run_browser_e2e() -> None:
    runtime_root = Path(".tmp") / "browser_e2e_runtime"
    if runtime_root.exists():
        shutil.rmtree(runtime_root)
    runtime_root.mkdir(parents=True, exist_ok=True)

    isolated_note_files_dir = configure_isolated_app(runtime_root)
    fixtures = create_docx_fixtures(runtime_root)

    port = reserve_port()
    base_url = f"http://127.0.0.1:{port}"

    config = uvicorn.Config(main_module.app, host="127.0.0.1", port=port, log_level="warning")
    server = uvicorn.Server(config)
    thread = threading.Thread(target=server.run, daemon=True)
    thread.start()

    try:
        wait_for_server_ready(base_url)

        with sync_playwright() as playwright:
            browser = playwright.chromium.launch(headless=True)
            page = browser.new_page()
            page.set_default_timeout(20_000)

            try:
                register_and_login(page, base_url, "e2e_topic_001", "password123")
                create_schedule_and_course(page)
                open_knowledge_workspace(page)

                create_topic(page, TOPIC_LOGIC, "命题、真值表、等值演算与范式", "真值表, 等值演算, 范式")
                create_topic(page, TOPIC_RELATION, "二元关系、等价关系、偏序与哈斯图", "等价关系, 偏序, 哈斯图")
                create_topic(page, TOPIC_GRAPH, "图、路径、欧拉回路与生成树", "图, 连通, 欧拉回路")

                upload_note(page, fixtures["logic"], "logic.docx")
                assign_latest_note_to_topic(page, TOPIC_LOGIC, "logic.docx")

                upload_note(page, fixtures["relation"], "relation.docx")
                assign_latest_note_to_topic(page, TOPIC_RELATION, "relation.docx")

                upload_note(page, fixtures["graph"], "graph.docx")
                assign_latest_note_to_topic(page, TOPIC_GRAPH, "graph.docx")

                show_all_notes(page)
                expect(page.locator(".note-card")).to_have_count(3)
                verify_search_and_graph(page)

                uploaded_files = list(isolated_note_files_dir.glob("*.docx"))
                assert len(uploaded_files) == 3, "Expected uploaded note files to land in the isolated runtime"
            finally:
                browser.close()
    finally:
        server.should_exit = True
        thread.join(timeout=10)


if __name__ == "__main__":
    run_browser_e2e()
