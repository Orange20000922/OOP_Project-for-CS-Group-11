from __future__ import annotations

import sys
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

import app.main as main_module
import app.routers.auth as auth_router
import app.routers.knowledge as knowledge_router
import app.routers.note as note_router
import app.routers.query as query_router
import app.routers.schedule as schedule_router
import app.services as services_module
import app.storage.schedule_store as schedule_store_module
import app.storage.user_store as user_store_module
from app.services.auth_service import AuthService
from app.services.knowledge_service import KnowledgeService
from app.services.note_service import NoteService
from app.services.schedule_service import ScheduleService
from app.storage.knowledge_tree_store import KnowledgeTreeStore
from app.storage.note_store import NoteStore
from app.storage.schedule_store import ScheduleStore
from app.storage.user_store import UserStore

note_service_module = sys.modules["app.services.note_service"]


class IntegrationScraperStub:
    def fetch_schedule(self, account, password, semester_id, *, prefer_playwright=False):
        return []

    def parse_pdf_schedule(self, content: bytes):
        return []


def make_docx_bytes(sections: list[tuple[str, list[str]]]) -> bytes:
    document = Document()
    for heading, paragraphs in sections:
        document.add_paragraph(heading)
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def note_knowledge_client(monkeypatch: pytest.MonkeyPatch, tmp_path: Path):
    data_dir = tmp_path / "data"
    users_file = data_dir / "users.json"
    schedules_dir = data_dir / "schedules"
    logs_dir = data_dir / "logs"
    note_files_dir = data_dir / "note_files"
    notes_db_path = data_dir / "notes.db"
    knowledge_tree_dir = data_dir / "knowledge_trees"

    monkeypatch.setattr(user_store_module, "USERS_FILE", users_file)
    monkeypatch.setattr(schedule_store_module, "SCHEDULES_DIR", schedules_dir)
    monkeypatch.setattr(note_service_module, "NOTE_FILES_DIR", note_files_dir)
    monkeypatch.setattr(main_module, "DATA_DIR", data_dir)
    monkeypatch.setattr(main_module, "USERS_FILE", users_file)
    monkeypatch.setattr(main_module, "SCHEDULES_DIR", schedules_dir)
    monkeypatch.setattr(main_module, "LOGS_DIR", logs_dir)
    monkeypatch.setattr(main_module, "APP_LOG_FILE", logs_dir / "app.log")
    monkeypatch.setattr(main_module, "NOTE_FILES_DIR", note_files_dir)

    user_store = UserStore()
    schedule_store = ScheduleStore()
    note_store = NoteStore(db_path=notes_db_path)
    scraper = IntegrationScraperStub()
    auth_service = AuthService(user_store)
    schedule_service = ScheduleService(schedule_store, scraper)
    note_service = NoteService(note_store)
    knowledge_service = KnowledgeService(
        note_store,
        tree_store=KnowledgeTreeStore(root_dir=knowledge_tree_dir),
    )
    knowledge_service._memory_failed = True
    knowledge_service._topic_store_failed = True
    knowledge_service._llm_failed = True

    monkeypatch.setattr(services_module, "user_store", user_store)
    monkeypatch.setattr(services_module, "schedule_store", schedule_store)
    monkeypatch.setattr(services_module, "note_store", note_store)
    monkeypatch.setattr(services_module, "scnu_scraper", scraper)
    monkeypatch.setattr(services_module, "auth_service", auth_service)
    monkeypatch.setattr(services_module, "schedule_service", schedule_service)
    monkeypatch.setattr(services_module, "note_service", note_service)
    monkeypatch.setattr(services_module, "knowledge_service", knowledge_service)
    monkeypatch.setattr(main_module, "auth_service", auth_service)

    monkeypatch.setattr(auth_router, "auth_service", auth_service)
    monkeypatch.setattr(schedule_router, "auth_service", auth_service)
    monkeypatch.setattr(schedule_router, "schedule_service", schedule_service)
    monkeypatch.setattr(query_router, "auth_service", auth_service)
    monkeypatch.setattr(query_router, "schedule_service", schedule_service)
    monkeypatch.setattr(note_router, "auth_service", auth_service)
    monkeypatch.setattr(note_router, "note_service", note_service)
    monkeypatch.setattr(note_router, "knowledge_service", knowledge_service)
    monkeypatch.setattr(knowledge_router, "auth_service", auth_service)
    monkeypatch.setattr(knowledge_router, "knowledge_service", knowledge_service)

    with TestClient(main_module.app) as client:
        yield client


def test_discrete_math_note_topic_integration(note_knowledge_client: TestClient):
    client = note_knowledge_client

    register_response = client.post(
        "/auth/register",
        json={
            "student_id": "20256001",
            "name": "离散数学测试",
            "password": "password123",
            "scnu_account": "20256001",
        },
    )
    assert register_response.status_code == 201

    login_response = client.post(
        "/auth/login",
        json={"student_id": "20256001", "password": "password123"},
    )
    assert login_response.status_code == 200

    init_response = client.post(
        "/schedule",
        json={"semester": "2025-2026-2", "semester_start": "2026-03-02"},
    )
    assert init_response.status_code == 201

    course_response = client.post(
        "/schedule/course",
        json={
            "name": "离散数学",
            "teacher": "周老师",
            "location": "理工楼 B201",
            "weekday": 2,
            "period_start": 3,
            "period_end": 4,
            "weeks": [1, 2, 3, 4, 5, 6],
            "week_type": "all",
        },
    )
    assert course_response.status_code == 201
    course_id = course_response.json()["id"]

    topic_ids: dict[str, str] = {}
    topic_payloads = [
        {
            "name": "命题逻辑",
            "summary": "命题、真值表、等值演算与范式",
            "keywords": ["真值表", "等值演算", "范式"],
        },
        {
            "name": "关系与偏序",
            "summary": "关系性质、等价关系、偏序与哈斯图",
            "keywords": ["等价关系", "偏序", "哈斯图"],
        },
        {
            "name": "图论",
            "summary": "图、路径、连通性、欧拉回路与生成树",
            "keywords": ["图", "连通", "欧拉回路"],
        },
    ]

    for payload in topic_payloads:
        response = client.post(
            "/knowledge/tree/topic",
            json={
                "course_id": course_id,
                "name": payload["name"],
                "summary": payload["summary"],
                "keywords": payload["keywords"],
                "parent_id": None,
            },
        )
        assert response.status_code == 200
        tree = response.json()
        for topic_id, topic in tree["topics"].items():
            if topic["name"] == payload["name"]:
                topic_ids[payload["name"]] = topic_id
                break

    assert set(topic_ids) == {"命题逻辑", "关系与偏序", "图论"}

    note_files = {
        "logic.docx": make_docx_bytes(
            [
                (
                    "1. 命题逻辑与等值演算",
                    [
                        "命题逻辑研究命题、联结词、真值表与等值演算。真值表可以判断命题公式的真假。",
                    ],
                ),
                (
                    "2. 范式与推理",
                    [
                        "主析取范式和主合取范式用于化简命题公式。利用等值演算可以判断推理是否有效。",
                    ],
                ),
            ]
        ),
        "relation.docx": make_docx_bytes(
            [
                (
                    "1. 二元关系",
                    [
                        "关系的基本性质包括自反性、对称性、反对称性和传递性。等价关系会把集合划分为等价类。",
                    ],
                ),
                (
                    "2. 偏序与哈斯图",
                    [
                        "偏序关系满足自反、反对称和传递。哈斯图可以展示整除关系与集合包含关系。",
                    ],
                ),
            ]
        ),
        "graph.docx": make_docx_bytes(
            [
                (
                    "1. 图与连通性",
                    [
                        "图论研究顶点、边、路径与连通分量。无向图的度数和连通性用于分析结构。",
                    ],
                ),
                (
                    "2. 欧拉回路与生成树",
                    [
                        "欧拉回路要求图连通且所有顶点度数为偶数。生成树保持连通而不产生回路。",
                    ],
                ),
            ]
        ),
    }

    uploaded_note_ids: dict[str, str] = {}
    for filename, content in note_files.items():
        upload_response = client.post(
            f"/note/upload?course_id={course_id}",
            files={
                "file": (
                    filename,
                    content,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert upload_response.status_code == 201
        detail = upload_response.json()
        assert detail["note"]["filename"] == filename
        assert detail["note"]["chunk_count"] == 2
        uploaded_note_ids[filename] = detail["note"]["id"]

    for topic_name, filename in {
        "命题逻辑": "logic.docx",
        "关系与偏序": "relation.docx",
        "图论": "graph.docx",
    }.items():
        assign_response = client.post(
            f"/knowledge/tree/topic/{topic_ids[topic_name]}/assign",
            json={
                "course_id": course_id,
                "note_id": uploaded_note_ids[filename],
            },
        )
        assert assign_response.status_code == 200

    notes_response = client.get("/note/list", params={"course_id": course_id})
    assert notes_response.status_code == 200
    assert len(notes_response.json()) == 3

    tree_response = client.get("/knowledge/tree", params={"course_id": course_id})
    assert tree_response.status_code == 200
    tree = tree_response.json()
    assert tree["topics"][topic_ids["命题逻辑"]]["note_ids"] == [uploaded_note_ids["logic.docx"]]
    assert tree["topics"][topic_ids["关系与偏序"]]["note_ids"] == [uploaded_note_ids["relation.docx"]]
    assert tree["topics"][topic_ids["图论"]]["note_ids"] == [uploaded_note_ids["graph.docx"]]

    search_response = client.post(
        "/knowledge/search",
        json={
            "query": "真值表",
            "limit": 5,
            "course_id": course_id,
        },
    )
    assert search_response.status_code == 200
    search_results = search_response.json()
    assert len(search_results) >= 1
    assert search_results[0]["chunk"]["note_id"] == uploaded_note_ids["logic.docx"]
    assert "命题逻辑" in search_results[0]["note_title"]

    graph_response = client.get(
        "/knowledge/graph",
        params={
            "course_id": course_id,
            "query": "欧拉回路",
            "top_k": 2,
            "min_score": 0.1,
            "max_nodes": 20,
            "topic_limit": 1,
        },
    )
    assert graph_response.status_code == 200
    graph = graph_response.json()
    assert graph["routing_applied"] is True
    assert graph["selected_topic_ids"] == [topic_ids["图论"]]
    assert {node["note_id"] for node in graph["nodes"]} == {uploaded_note_ids["graph.docx"]}
    assert graph["total_nodes"] == 2
    assert graph["total_links"] >= 1
